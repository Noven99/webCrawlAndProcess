#!/usr/bin/env python
# fetch_ai_news.py
"""
抓取 Amnesty International “News” 文章正文并保存为 A4 Word 文档
"""

import argparse
import re
import sys
from pathlib import Path
from typing import List

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/126.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
}

YAHEI = "Microsoft YaHei"  # 用于中文显示（可按需修改）


# ---------------- 网络抓取 ----------------
def fetch_page(url: str) -> str:
    resp = requests.get(url, headers=HEADERS, timeout=15)
    resp.raise_for_status()
    return resp.text


# ---------------- 解析 ----------------
def parse_article(html: str) -> dict:
    soup = BeautifulSoup(html, "lxml")

    # 标题
    title_tag = soup.find("h1")
    title = title_tag.get_text(strip=True) if title_tag else "Untitled"

    # 日期
    pub = ""
    time_tag = soup.find("time")
    if time_tag and time_tag.has_attr("datetime"):
        pub = time_tag["datetime"][:10]  # YYYY-MM-DD
    else:
        m = re.search(r"\b\d{1,2}\s\w+\s\d{4}\b", soup.get_text())
        if m:
            pub = m.group(0)

    # 正文
    container = (
            soup.find("div", class_=re.compile(r"(rich-text|article-body)"))
            or soup.find("article")
    )
    paragraphs: List[str] = []
    if container:
        for p in container.find_all("p"):
            txt = p.get_text(" ", strip=True)
            if txt:
                paragraphs.append(txt)

    return {"title": title, "date": pub, "paragraphs": paragraphs}


# ---------------- Word 导出 ----------------
def save_to_word(article: dict, path: Path):
    """将爬取的内容保存到 Word 文件"""
    path.parent.mkdir(parents=True, exist_ok=True)
    # 创建 Word 文档
    doc = Document()

    # 设置页面为 A4 尺寸
    section = doc.sections[0]
    section.page_height = Inches(11.7)  # A4 高度
    section.page_width = Inches(8.3)   # A4 宽度
    section.left_margin = Inches(1)   # 左边距
    section.right_margin = Inches(1)  # 右边距
    section.top_margin = Inches(1)    # 上边距
    section.bottom_margin = Inches(1) # 下边距

    # 添加标题
    title = article["title"]
    doc.add_heading(title, level=1)

    # 添加发布日期
    if article["date"]:
        date_paragraph = doc.add_paragraph(f"发布日期：{article['date']}")
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 添加正文内容，每段独立一个段落
    for para in article["paragraphs"]:
        paragraph = doc.add_paragraph(para)
        paragraph.style.font.size = Pt(12)  # 设置字体大小为 12pt

    # 保存文件
    doc.save(path)
    print(f"已保存到 {path}")


# ---------------- CLI ----------------
def sanitize_filename(text: str) -> str:
    text = re.sub(r"[\\/:*?\"<>|]", "_", text)
    return re.sub(r"\s+", "_", text).strip("_")


def main():
    parser = argparse.ArgumentParser(
        description="Download Amnesty International news article(s) to A4 Word files."
    )
    parser.add_argument(
        "urls",
        nargs="+",
        help="文章 URL，可一次给多个",
    )
    parser.add_argument(
        "--outdir",
        default=".",
        help="导出目录（默认当前目录）",
    )
    args = parser.parse_args()

    out_dir = Path(args.outdir).expanduser().resolve()
    for url in args.urls:
        try:
            html = fetch_page(url)
            art = parse_article(html)
            fname = sanitize_filename(
                f"{art['date'] or 'undated'}-{art['title']}.docx"
            )
            save_to_word(art, out_dir / fname)
        except Exception as e:
            print(f"❌ 处理 {url} 时出错：{e}", file=sys.stderr)


if __name__ == "__main__":
    main()