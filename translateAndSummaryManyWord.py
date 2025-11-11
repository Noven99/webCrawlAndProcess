"""

1. 逐段翻译英文→简体中文（译文紧跟原文，字体“微软雅黑”）
2. 若英文段落单词数 > 40，则追加“段落总结”（微软雅黑、加粗）

"""

import sys
from pathlib import Path
from typing import List
from docx.shared import Pt
from docx import Document
from docx.oxml.ns import qn

# ---------------- NLTK ------------------------------------------------------
import nltk

nltk_data = Path(__file__).parent / "nltk_data"
nltk.data.path.append(str(nltk_data))
for pkg in ["punkt", "stopwords"]:
    try:
        nltk.data.find(f"corpora/{pkg}" if pkg == "stopwords" else f"tokenizers/{pkg}")
    except LookupError:
        nltk.download(pkg, download_dir=str(nltk_data))

from nltk.tokenize import word_tokenize, sent_tokenize

# ---------------- 翻译 -------------------------------------------------------
from deep_translator import GoogleTranslator

_trans = GoogleTranslator(source="en", target="zh-CN")


def translate_en2zh(text: str) -> str:
    if not text.strip():
        return ""
    try:
        return _trans.translate(text)
    except Exception:
        return text


# ---------------- 摘要模型 ---------------------------------------------------
from transformers import pipeline, BartTokenizer

SUMMARIZER = pipeline(
    "summarization",
    model="facebook/bart-large-cnn",
    tokenizer="facebook/bart-large-cnn",
    framework="pt",
    device=0          #  GPU
)
TOKENIZER = BartTokenizer.from_pretrained("facebook/bart-large-cnn")

# ---------------- 参数 -------------------------------------------------------
MIN_WORDS = 40
MAX_INPUT_TOKENS = 1024          # BART 输入上限
MIN_LEN_FIXED = 30               # 摘要最小长度
MAX_LEN_FIXED = 80               # 摘要最大长度（固定 80）
BATCH_SIZE = 4                   # 显存不足可调小
YAHEI = "Microsoft YaHei"


# ---------------- 工具 -------------------------------------------------------
def truncate_text_by_sentences(text: str,
                               max_tokens: int = MAX_INPUT_TOKENS) -> str:
    """句子级截断，使 token 数 ≤ max_tokens"""
    sents = sent_tokenize(text)
    cur = 0
    kept = []
    for s in sents:
        tl = len(TOKENIZER.tokenize(s))
        if cur + tl > max_tokens:
            break
        kept.append(s)
        cur += tl
    return " ".join(kept) + "..." if kept else text


class Candidate:
    """待摘要段落的信息"""
    def __init__(self, para_idx: int, text: str):
        self.idx = para_idx
        self.truncated = (truncate_text_by_sentences(text)
                          if len(TOKENIZER.tokenize(text)) > MAX_INPUT_TOKENS
                          else text)
        self.summary_zh = ""


def collect_candidates(paragraphs) -> List[Candidate]:
    """找出需要摘要的段落（单词数 > MIN_WORDS）"""
    cands = []
    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        if len([w for w in word_tokenize(txt) if w.isalpha()]) > MIN_WORDS:
            cands.append(Candidate(i, txt))
    return cands


def batch_summarize(cands: List[Candidate]) -> None:
    """GPU 批量生成摘要，写回 cand.summary_zh"""
    if not cands:
        return

    texts = [c.truncated for c in cands]
    for start in range(0, len(texts), BATCH_SIZE):
        chunk = texts[start:start + BATCH_SIZE]
        outs = SUMMARIZER(
            chunk,
            max_length=MAX_LEN_FIXED,
            min_length=MIN_LEN_FIXED,
            do_sample=False,
            truncation=True,
            batch_size=len(chunk)        # 让 pipeline 一次性吃掉此批
        )
        for c, out in zip(cands[start:start + BATCH_SIZE], outs):
            c.summary_zh = translate_en2zh(out["summary_text"].strip())


# ---------------- Word 生成后的格式 --------------------------------------------------
def set_yahei(run, bold=False, size: int = 11):
    run.font.name = YAHEI
    run._element.rPr.rFonts.set(qn('w:eastAsia'), YAHEI)
    run.font.size = Pt(size)
    run.bold = bold


# ---------------- 主流程 -----------------------------------------------------
def process_docx(path: Path):
    doc = Document(path)
    total_para = len(doc.paragraphs)

    # 1. 收集 & 批量摘要
    cands = collect_candidates(doc.paragraphs)
    print(f"共有 {len(cands)} 个有效段落需要进行总结。")
    batch_summarize(cands)
    idx2summary = {c.idx: c.summary_zh for c in cands}

    # 2. 遍历段落，追加翻译 / 摘要
    for i, para in enumerate(doc.paragraphs):
        print(f"正在处理第 word 文件的第 {i + 1}/{total_para} 段")
        txt = para.text.strip()
        if not txt:
            continue

        # 2-1 中文翻译
        zh = translate_en2zh(txt)
        if zh:
            r_tr = para.add_run("\n" + zh)
            set_yahei(r_tr, False)

        # 2-2 段落总结
        if i in idx2summary and idx2summary[i]:
            r_sm = para.add_run("\n段落总结：" + idx2summary[i])
            set_yahei(r_sm, True)

    # 3. 保存
    out = path.with_name(f"{path.stem}_zh.docx")
    doc.save(out)
    print("处理完成 →", out.resolve())


# ---------------- CLI -------------------------------------------------------
if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(
        description="Translate English paragraphs to Chinese and add summaries to long ones."
    )
    parser.add_argument(
        "files",
        nargs="+",
        help="要处理的 .docx 文件，可以一次给多个，也支持通配符 *.docx"
    )
    args = parser.parse_args()

    for f in args.files:
        file_in = Path(f)
        if not file_in.exists():
            print("❌  文件不存在：", file_in)
            continue
        if file_in.suffix.lower() != ".docx":
            print("⚠️  忽略非 .docx 文件：", file_in)
            continue
        process_docx(file_in)